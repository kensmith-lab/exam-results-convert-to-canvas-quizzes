#!/usr/bin/env python3
"""
Enhanced Canvas QTI Quiz Generator - Multi-Format Support
Processes entire folder structures with various document types:
- HTML, PDF, DOCX, XLSX, TXT files
- Batch processing of nested folder structures
- Enhanced question pattern recognition for AWS exam formats
"""

import os
import re
import sys
import uuid
import zipfile
import tempfile
import xml.etree.ElementTree as ET
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any, Tuple, Optional
import glob

# Import required libraries with auto-install
def install_and_import(package):
    try:
        __import__(package)
    except ImportError:
        print(f"Installing {package}...")
        os.system(f"pip install {package}")

# Install required packages
packages = ['beautifulsoup4', 'PyPDF2', 'python-docx', 'lxml', 'pandas', 'openpyxl', 'xlrd']
for package in packages:
    if package == 'beautifulsoup4':
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            install_and_import('beautifulsoup4')
            from bs4 import BeautifulSoup
    elif package == 'PyPDF2':
        try:
            import PyPDF2
        except ImportError:
            install_and_import('PyPDF2')
            import PyPDF2
    elif package == 'python-docx':
        try:
            from docx import Document
        except ImportError:
            install_and_import('python-docx')
            from docx import Document
    elif package == 'pandas':
        try:
            import pandas as pd
        except ImportError:
            install_and_import('pandas')
            import pandas as pd
    elif package == 'openpyxl':
        try:
            import openpyxl
        except ImportError:
            install_and_import('openpyxl')
            import openpyxl


class Question:
    """Represents a quiz question"""
    def __init__(self, question_text: str, question_type: str, choices: List[str], 
                 correct_answers: List[int], points: int = 1, source_file: str = ""):
        self.id = str(uuid.uuid4())
        self.question_text = question_text.strip()
        self.question_type = question_type  # 'multiple_choice' or 'multiple_select'
        self.choices = [choice.strip() for choice in choices]
        self.correct_answers = correct_answers  # List of indices (0-based)
        self.points = points
        self.source_file = source_file


class EnhancedDocumentParser:
    """Enhanced parser for multiple document types and folder structures"""
    
    @staticmethod
    def process_folder(folder_path: str) -> List[Question]:
        """Process entire folder structure recursively"""
        folder_path = Path(folder_path)
        if not folder_path.exists():
            print(f"❌ Folder not found: {folder_path}")
            return []
        
        all_questions = []
        supported_extensions = ['.html', '.htm', '.pdf', '.docx', '.xlsx', '.xls', '.txt']
        
        print(f"🔍 Scanning folder: {folder_path}")
        
        # Find all supported files recursively
        for ext in supported_extensions:
            pattern = f"**/*{ext}"
            files = list(folder_path.glob(pattern))
            
            for file_path in files:
                print(f"📄 Processing: {file_path.relative_to(folder_path)}")
                
                try:
                    questions = EnhancedDocumentParser.parse_file(file_path)
                    if questions:
                        all_questions.extend(questions)
                        print(f"   ✅ Found {len(questions)} questions")
                    else:
                        print(f"   ⚠️  No questions found")
                        
                except Exception as e:
                    print(f"   ❌ Error processing {file_path.name}: {e}")
        
        return all_questions
    
    @staticmethod
    def parse_file(file_path: Path) -> List[Question]:
        """Parse any supported file type"""
        questions = []
        
        try:
            if file_path.suffix.lower() in ['.html', '.htm']:
                questions = EnhancedDocumentParser.parse_html(str(file_path))
            elif file_path.suffix.lower() == '.pdf':
                questions = EnhancedDocumentParser.parse_pdf(str(file_path))
            elif file_path.suffix.lower() == '.docx':
                questions = EnhancedDocumentParser.parse_docx(str(file_path))
            elif file_path.suffix.lower() in ['.xlsx', '.xls']:
                questions = EnhancedDocumentParser.parse_excel(str(file_path))
            elif file_path.suffix.lower() == '.txt':
                questions = EnhancedDocumentParser.parse_text_file(str(file_path))
            else:
                print(f"⚠️  Unsupported file format: {file_path.suffix}")
            
            # Add source file info to each question
            for question in questions:
                question.source_file = str(file_path.name)
                
        except Exception as e:
            print(f"❌ Error parsing {file_path.name}: {e}")
        
        return questions
    
    @staticmethod
    def parse_excel(file_path: str) -> List[Question]:
        """Parse Excel/XLSX files for questions"""
        questions = []
        
        try:
            # Try reading with pandas first
            df = pd.read_excel(file_path, sheet_name=None)  # Read all sheets
            
            for sheet_name, sheet_df in df.items():
                print(f"   📊 Processing sheet: {sheet_name}")
                
                # Convert dataframe to text content
                text_content = ""
                for _, row in sheet_df.iterrows():
                    row_text = " ".join([str(cell) for cell in row if pd.notna(cell)])
                    if row_text.strip():
                        text_content += row_text + "\n"
                
                # Parse the text content
                sheet_questions = EnhancedDocumentParser._parse_text_content(text_content)
                questions.extend(sheet_questions)
                
        except Exception as e:
            print(f"   ⚠️  Excel parsing error: {e}")
        
        return questions
    
    @staticmethod
    def parse_text_file(file_path: str) -> List[Question]:
        """Parse plain text files"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text_content = f.read()
        return EnhancedDocumentParser._parse_text_content(text_content)
    
    @staticmethod
    def parse_html(file_path: str) -> List[Question]:
        """Parse HTML file for questions"""
        questions = []
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            soup = BeautifulSoup(f.read(), 'html.parser')
            
        # Look for common question patterns
        question_blocks = soup.find_all(['div', 'p', 'section', 'li', 'td'], 
                                      class_=re.compile(r'question|quiz|item|test', re.I))
        
        if not question_blocks:
            # Fallback: parse text content directly
            text_content = soup.get_text()
            questions = EnhancedDocumentParser._parse_text_content(text_content)
        else:
            for block in question_blocks:
                question = EnhancedDocumentParser._extract_question_from_text(block.get_text())
                if question:
                    questions.append(question)
        
        return questions
    
    @staticmethod
    def parse_pdf(file_path: str) -> List[Question]:
        """Enhanced PDF parsing for exam questions"""
        questions = []
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                text_content = ""
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        text_content += f"\n--- Page {page_num + 1} ---\n" + page_text + "\n"
                    except Exception as e:
                        print(f"   ⚠️  Error reading page {page_num + 1}: {e}")
                        continue
            
            questions = EnhancedDocumentParser._parse_text_content(text_content)
            
        except Exception as e:
            print(f"   ❌ PDF parsing error: {e}")
        
        return questions
    
    @staticmethod
    def parse_docx(file_path: str) -> List[Question]:
        """Parse DOCX file for questions"""
        try:
            doc = Document(file_path)
            text_content = ""
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content += para.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        text_content += row_text + "\n"
            
            return EnhancedDocumentParser._parse_text_content(text_content)
            
        except Exception as e:
            print(f"   ❌ DOCX parsing error: {e}")
            return []
    
    @staticmethod
    def _parse_text_content(text: str) -> List[Question]:
        """Enhanced text parsing with AWS exam question patterns"""
        questions = []
        
        # Clean up text
        text = re.sub(r'\s+', ' ', text)  # Normalize whitespace
        text = re.sub(r'Page \d+', '', text)  # Remove page markers
        
        # AWS exam specific patterns
        aws_patterns = [
            # Pattern 1: Question X: ... A) B) C) D)
            r'(?:Question\s*\d+[:.]\s*)(.*?)(?=Question\s*\d+[:.]\s*|\n\s*Answer[:\s]|\Z)',
            # Pattern 2: Q\d+. ... A. B. C. D.
            r'(?:Q\d+\.\s*)(.*?)(?=Q\d+\.\s*|\n\s*Answer[:\s]|\Z)',
            # Pattern 3: \d+\. ... A\) B\) C\) D\)
            r'(?:^\d+\.\s*)(.*?)(?=^\d+\.\s*|\n\s*Answer[:\s]|\Z)',
        ]
        
        # Try AWS-specific patterns first
        for pattern in aws_patterns:
            matches = re.finditer(pattern, text, re.MULTILINE | re.DOTALL)
            for match in matches:
                question_text = match.group(1).strip()
                if len(question_text) > 20:  # Minimum question length
                    question = EnhancedDocumentParser._extract_question_from_aws_text(question_text)
                    if question:
                        questions.append(question)
        
        # If no AWS patterns found, try general patterns
        if not questions:
            questions = EnhancedDocumentParser._parse_general_text_content(text)
        
        return questions
    
    @staticmethod
    def _extract_question_from_aws_text(text: str) -> Optional[Question]:
        """Extract question from AWS exam format text"""
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        
        if len(lines) < 3:  # Need at least question + 2 choices
            return None
        
        # Find the main question (usually the longest line or contains ?)
        question_text = ""
        choice_start_idx = 0
        
        for i, line in enumerate(lines):
            # Look for question indicators
            if ('?' in line or 
                len(line) > 50 or  # Likely the main question if long
                re.search(r'which|what|how|when|where|why|should|would|best|most', line, re.I)):
                question_text = line
                choice_start_idx = i + 1
                break
        
        if not question_text:
            question_text = lines[0]
            choice_start_idx = 1
        
        # Extract choices with various patterns
        choices = []
        choice_patterns = [
            r'^[A-Z]\)\s*(.+)',      # A) choice
            r'^[A-Z]\.\s*(.+)',      # A. choice
            r'^\([A-Z]\)\s*(.+)',    # (A) choice
            r'^[A-Z]:\s*(.+)',       # A: choice
        ]
        
        for line in lines[choice_start_idx:]:
            for pattern in choice_patterns:
                match = re.match(pattern, line)
                if match:
                    choice_text = match.group(1).strip()
                    if choice_text:
                        choices.append(choice_text)
                    break
        
        if len(choices) < 2:
            return None
        
        # Look for correct answers (various patterns)
        correct_answers = []
        correct_patterns = [
            r'\*|correct|right|answer',
            r'✓|✔',
            r'\[correct\]|\(correct\)',
        ]
        
        for i, choice in enumerate(choices):
            for pattern in correct_patterns:
                if re.search(pattern, choice, re.I):
                    correct_answers.append(i)
                    choices[i] = re.sub(pattern, '', choice, flags=re.I).strip()
                    break
        
        # Default to first choice if no correct answer found
        if not correct_answers:
            correct_answers = [0]
        
        # Determine question type
        question_type = 'multiple_select' if len(correct_answers) > 1 else 'multiple_choice'
        
        return Question(question_text, question_type, choices, correct_answers)
    
    @staticmethod
    def _parse_general_text_content(text: str) -> List[Question]:
        """General text parsing (original method)"""
        questions = []
        
        # Split text into potential question blocks
        blocks = re.split(r'\n\s*\n', text)
        
        for block in blocks:
            if EnhancedDocumentParser._looks_like_question(block):
                question = EnhancedDocumentParser._extract_question_from_text(block)
                if question:
                    questions.append(question)
        
        return questions
    
    @staticmethod
    def _looks_like_question(text: str) -> bool:
        """Check if text block looks like a question"""
        text = text.strip()
        if len(text) < 10:
            return False
        
        # Check for question markers
        has_question_marker = bool(re.search(r'(?:question\s*\d*[:.?]|\?)', text, re.I))
        has_choices = bool(re.search(r'[A-Za-z]\)|\d+\.', text))
        
        return has_question_marker or has_choices
    
    @staticmethod
    def _extract_question_from_text(text: str) -> Optional[Question]:
        """Extract question from general text block"""
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        
        if len(lines) < 3:  # Need at least question + 2 choices
            return None
        
        # Find question text (usually first line or contains ?)
        question_text = ""
        choice_start_idx = 0
        
        for i, line in enumerate(lines):
            if '?' in line or re.match(r'question\s*\d*[:.]\s*', line, re.I):
                question_text = re.sub(r'question\s*\d*[:.]\s*', '', line, flags=re.I).strip()
                choice_start_idx = i + 1
                break
        
        if not question_text and lines:
            question_text = lines[0]
            choice_start_idx = 1
        
        # Extract choices
        choices = []
        for line in lines[choice_start_idx:]:
            # Remove choice markers (A), 1., etc.
            clean_choice = re.sub(r'^[A-Za-z]\)|\d+\.\s*', '', line).strip()
            if clean_choice:
                choices.append(clean_choice)
        
        if len(choices) < 2:
            return None
        
        # Determine correct answers (look for markers like *, CORRECT, etc.)
        correct_answers = []
        for i, choice in enumerate(choices):
            if re.search(r'\*|correct|right|answer', choice, re.I):
                correct_answers.append(i)
                choices[i] = re.sub(r'\s*\*|correct|right|answer', '', choice, flags=re.I).strip()
        
        # Default to first choice if no correct answer found
        if not correct_answers:
            correct_answers = [0]
        
        # Determine question type
        question_type = 'multiple_select' if len(correct_answers) > 1 else 'multiple_choice'
        
        return Question(question_text, question_type, choices, correct_answers)


class QTIGenerator:
    """Generates QTI-compliant XML and zip files"""
    
    def __init__(self, quiz_title: str = "Multi-Format Quiz Import"):
        self.quiz_title = quiz_title
        self.quiz_id = str(uuid.uuid4())
        
    def generate_qti_zip(self, questions: List[Question], output_path: str):
        """Generate complete QTI zip file with source file info"""
        
        # Create temporary directory structure
        temp_dir = Path("temp_qti")
        temp_dir.mkdir(exist_ok=True)
        
        try:
            # Generate manifest file
            self._create_manifest(temp_dir, questions)
            
            # Generate assessment XML
            self._create_assessment_xml(temp_dir, questions)
            
            # Create zip file
            with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for file_path in temp_dir.rglob('*'):
                    if file_path.is_file():
                        arcname = str(file_path.relative_to(temp_dir))
                        zipf.write(file_path, arcname)
            
            print(f"📦 QTI zip file created: {output_path}")
            
        finally:
            # Cleanup
            import shutil
            shutil.rmtree(temp_dir, ignore_errors=True)
    
    def _create_manifest(self, temp_dir: Path, questions: List[Question]):
        """Create imsmanifest.xml"""
        
        manifest = ET.Element("manifest")
        manifest.set("xmlns", "http://www.imsglobal.org/xsd/imsccv1p1/imscp_v1p1")
        manifest.set("xmlns:lom", "http://ltsc.ieee.org/xsd/imsccv1p1/LOM/resource")
        manifest.set("xmlns:imsqti", "http://www.imsglobal.org/xsd/imsqti_v2p1")
        manifest.set("identifier", f"man_{self.quiz_id}")
        
        # Metadata
        metadata = ET.SubElement(manifest, "metadata")
        lom_general = ET.SubElement(metadata, "lom:general")
        lom_title = ET.SubElement(lom_general, "lom:title")
        lom_string = ET.SubElement(lom_title, "lom:string")
        lom_string.text = self.quiz_title
        
        # Organizations
        organizations = ET.SubElement(manifest, "organizations")
        organizations.set("default", f"org_{self.quiz_id}")
        
        organization = ET.SubElement(organizations, "organization")
        organization.set("identifier", f"org_{self.quiz_id}")
        
        title = ET.SubElement(organization, "title")
        title.text = self.quiz_title
        
        # Resources
        resources = ET.SubElement(manifest, "resources")
        
        # Assessment resource
        assessment_resource = ET.SubElement(resources, "resource")
        assessment_resource.set("identifier", f"assessment_{self.quiz_id}")
        assessment_resource.set("type", "imsqti_xmlv2p1")
        assessment_resource.set("href", "assessment.xml")
        
        file_elem = ET.SubElement(assessment_resource, "file")
        file_elem.set("href", "assessment.xml")
        
        # Write manifest
        tree = ET.ElementTree(manifest)
        ET.indent(tree, space="  ", level=0)
        tree.write(temp_dir / "imsmanifest.xml", encoding="utf-8", xml_declaration=True)
    
    def _create_assessment_xml(self, temp_dir: Path, questions: List[Question]):
        """Create assessment.xml with questions"""
        
        assessment = ET.Element("assessmentTest")
        assessment.set("xmlns", "http://www.imsglobal.org/xsd/imsqti_v2p1")
        assessment.set("identifier", f"assessment_{self.quiz_id}")
        assessment.set("title", self.quiz_title)
        
        # Test part
        test_part = ET.SubElement(assessment, "testPart")
        test_part.set("identifier", "testpart_1")
        test_part.set("navigationMode", "linear")
        test_part.set("submissionMode", "individual")
        
        # Assessment section
        section = ET.SubElement(test_part, "assessmentSection")
        section.set("identifier", "section_1")
        section.set("title", "Questions")
        section.set("visible", "true")
        
        # Add questions
        for i, question in enumerate(questions):
            item_ref = ET.SubElement(section, "assessmentItemRef")
            item_ref.set("identifier", f"item_{question.id}")
            item_ref.set("href", f"item_{question.id}.xml")
            
            # Create individual question file
            self._create_question_xml(temp_dir, question)
        
        # Write assessment
        tree = ET.ElementTree(assessment)
        ET.indent(tree, space="  ", level=0)
        tree.write(temp_dir / "assessment.xml", encoding="utf-8", xml_declaration=True)
    
    def _create_question_xml(self, temp_dir: Path, question: Question):
        """Create individual question XML file"""
        
        item = ET.Element("assessmentItem")
        item.set("xmlns", "http://www.imsglobal.org/xsd/imsqti_v2p1")
        item.set("identifier", f"item_{question.id}")
        item.set("title", question.question_text[:50] + "...")
        item.set("adaptive", "false")
        item.set("timeDependent", "false")
        
        # Response declaration
        response_decl = ET.SubElement(item, "responseDeclaration")
        response_decl.set("identifier", "RESPONSE")
        response_decl.set("cardinality", "multiple" if question.question_type == "multiple_select" else "single")
        response_decl.set("baseType", "identifier")
        
        # Correct response
        correct_response = ET.SubElement(response_decl, "correctResponse")
        for correct_idx in question.correct_answers:
            value = ET.SubElement(correct_response, "value")
            value.text = f"choice_{correct_idx}"
        
        # Outcome declaration
        outcome_decl = ET.SubElement(item, "outcomeDeclaration")
        outcome_decl.set("identifier", "SCORE")
        outcome_decl.set("cardinality", "single")
        outcome_decl.set("baseType", "float")
        
        default_value = ET.SubElement(outcome_decl, "defaultValue")
        value = ET.SubElement(default_value, "value")
        value.text = "0"
        
        # Item body
        item_body = ET.SubElement(item, "itemBody")
        
        # Question text with source file info
        div = ET.SubElement(item_body, "div")
        p = ET.SubElement(div, "p")
        p.text = question.question_text
        
        if question.source_file:
            source_p = ET.SubElement(div, "p")
            source_p.set("style", "font-size: small; color: gray;")
            source_p.text = f"Source: {question.source_file}"
        
        # Choice interaction
        interaction = ET.SubElement(item_body, "choiceInteraction")
        interaction.set("responseIdentifier", "RESPONSE")
        interaction.set("shuffle", "false")
        
        if question.question_type == "multiple_select":
            interaction.set("maxChoices", str(len(question.choices)))
        else:
            interaction.set("maxChoices", "1")
        
        # Add choices
        for i, choice_text in enumerate(question.choices):
            simple_choice = ET.SubElement(interaction, "simpleChoice")
            simple_choice.set("identifier", f"choice_{i}")
            simple_choice.text = choice_text
        
        # Response processing
        response_processing = ET.SubElement(item, "responseProcessing")
        response_processing.set("template", "http://www.imsglobal.org/question/qti_v2p1/rptemplates/match_correct")
        
        # Write question file
        tree = ET.ElementTree(item)
        ET.indent(tree, space="  ", level=0)
        tree.write(temp_dir / f"item_{question.id}.xml", encoding="utf-8", xml_declaration=True)


def main():
    """Main function with folder processing support"""
    
    print("🚀 Enhanced Canvas QTI Generator - Multi-Format Support")
    print("=====================================================")
    
    # Check for input folder argument
    if len(sys.argv) > 1:
        input_path = Path(sys.argv[1])
        if input_path.exists():
            print(f"📁 Processing folder: {input_path}")
            all_questions = EnhancedDocumentParser.process_folder(str(input_path))
        else:
            print(f"❌ Folder not found: {input_path}")
            return
    else:
        # Default behavior - process documents folder
        documents_dir = Path("documents")
        output_dir = Path("output")
        output_dir.mkdir(exist_ok=True)
        
        if not documents_dir.exists():
            print(f"Creating {documents_dir} directory...")
            documents_dir.mkdir(exist_ok=True)
            print(f"Please place your files in the '{documents_dir}' directory and run again.")
            return
        
        print(f"📁 Processing documents folder: {documents_dir}")
        all_questions = EnhancedDocumentParser.process_folder(str(documents_dir))
    
    if not all_questions:
        print("\n❌ No questions found in any documents.")
        print("\nSupported formats:")
        print("- HTML (.html, .htm)")
        print("- PDF (.pdf)")  
        print("- Word (.docx)")
        print("- Excel (.xlsx, .xls)")
        print("- Text (.txt)")
        print("\nExpected format:")
        print("Question: What is AWS?")
        print("A) Amazon Web Services *")
        print("B) A cloud platform")
        print("C) A database")
        return
    
    # Group questions by source file
    source_files = {}
    for question in all_questions:
        source = question.source_file or "unknown"
        if source not in source_files:
            source_files[source] = 0
        source_files[source] += 1
    
    print(f"\n📊 Summary:")
    print(f"   Total questions found: {len(all_questions)}")
    print(f"   From {len(source_files)} files:")
    for source, count in source_files.items():
        print(f"     - {source}: {count} questions")
    
    # Generate QTI zip
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    # Create output directory if using folder argument
    if len(sys.argv) > 1:
        output_dir = Path("output")
        output_dir.mkdir(exist_ok=True)
    
    output_file = output_dir / f"multi_format_quiz_{timestamp}.zip"
    
    print(f"\n📦 Generating QTI zip file...")
    generator = QTIGenerator(f"Multi-Format Quiz Import {timestamp}")
    generator.generate_qti_zip(all_questions, str(output_file))
    
    print(f"\n✅ SUCCESS! QTI zip file ready for Canvas import:")
    print(f"   📁 {output_file}")
    print(f"   📝 {len(all_questions)} questions from multiple formats")
    
    print(f"\n🎯 Canvas Import Instructions:")
    print("1. Go to your Canvas course")
    print("2. Navigate to Quizzes")
    print("3. Click 'Import Quiz'")
    print("4. Select 'QTI .zip file'")
    print("5. Upload the generated .zip file")
    print("6. Review and publish your quiz")


if __name__ == "__main__":
    main()